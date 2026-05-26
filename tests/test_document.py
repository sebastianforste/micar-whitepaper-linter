from pathlib import Path
from unittest.mock import MagicMock, patch

import docx

from micar_linter.document import load_from_docx, load_from_pdf
from micar_linter.whitepaper import WhitepaperType, load_whitepaper


def test_load_from_docx(tmp_path: Path):
    doc_path = tmp_path / "test.docx"
    doc = docx.Document()
    
    # Add sections with standard headings and paragraphs
    doc.add_heading("Plain-Language Summary", level=1)
    doc.add_paragraph("This is the plain-language summary of the PayEUR token.")
    
    doc.add_heading("Information about the Issuer", level=1)
    doc.add_paragraph("PayEUR EMI GmbH is an authorised e-money institution.")
    doc.add_paragraph("Registered address is Kaiserstraße 12, 60311 Frankfurt am Main.")
    
    doc.add_heading("Risk Factors", level=1)
    doc.add_paragraph("The principal risks include market risk and credit risk.")
    
    doc.save(doc_path)
    
    sections = load_from_docx(doc_path)
    
    assert "summary" in sections
    assert "issuer" in sections
    assert "risks" in sections
    
    assert "This is the plain-language summary" in sections["summary"]
    assert "Frankfurt am Main" in sections["issuer"]
    assert "market risk" in sections["risks"]


@patch("pypdf.PdfReader")
def test_load_from_pdf(mock_reader_cls):
    mock_reader = MagicMock()
    mock_page1 = MagicMock()
    mock_page1.extract_text.return_value = "Zusammenfassung\nThis is a summary of the token."
    mock_page2 = MagicMock()
    mock_page2.extract_text.return_value = "Risikofaktoren\nThere are risk factors associated with this asset."
    
    mock_reader.pages = [mock_page1, mock_page2]
    mock_reader_cls.return_value = mock_reader
    
    sections = load_from_pdf(Path("dummy.pdf"))
    
    assert "summary" in sections
    assert "risks" in sections
    assert "This is a summary" in sections["summary"]
    assert "There are risk factors" in sections["risks"]


def test_load_whitepaper_docx_type_inference(tmp_path: Path):
    doc_path = tmp_path / "emt-test.docx"
    doc = docx.Document()
    
    doc.add_heading("Plain-Language Summary", level=1)
    doc.add_paragraph("This is the plain-language summary.")
    
    # Add keywords that trigger EMT type inference
    doc.add_heading("The E-Money Token", level=1)
    doc.add_paragraph("This e-money token has par-value redemption and safeguarding of funds.")
    
    doc.save(doc_path)
    
    wp = load_whitepaper(doc_path)
    assert wp.title == "Emt Test"
    assert wp.type == WhitepaperType.EMT
    assert "summary" in wp.sections
    assert "emt" in wp.sections


@patch("pypdf.PdfReader")
def test_load_whitepaper_pdf_type_inference(mock_reader_cls, tmp_path: Path):
    pdf_path = tmp_path / "art-test.pdf"
    
    mock_reader = MagicMock()
    mock_page1 = MagicMock()
    # Add keywords that trigger ART type inference
    mock_page1.extract_text.return_value = (
        "Zusammenfassung\nThis is a summary of the asset-referenced token.\n"
        "Reserve of assets stabilisation mechanism composition."
    )
    
    mock_reader.pages = [mock_page1]
    mock_reader_cls.return_value = mock_reader
    
    # Path must physically exist to pass the existence check in load_whitepaper,
    # so we create a blank file on disk and mock the actual reader reading it.
    pdf_path.write_bytes(b"")
    
    wp = load_whitepaper(pdf_path)
    assert wp.title == "Art Test"
    assert wp.type == WhitepaperType.ART
    assert "summary" in wp.sections
    assert "art" in wp.sections
